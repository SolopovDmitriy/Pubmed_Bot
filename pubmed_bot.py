import os
from datetime import datetime, timedelta
from telegram import Update
from telegram.ext import (
    ApplicationBuilder, CommandHandler, MessageHandler,
    filters, ContextTypes, ConversationHandler
)
from Bio import Entrez, Medline
import pandas as pd
from docx import Document

# === Вкажи свій email для PubMed ===
Entrez.email = "your_email@example.com"

# === Етапи діалогу ===
ASK_KEYWORDS, ASK_DAYS = range(2)

# === Пошук у PubMed ===
def fetch_pubmed_full(keywords, days=30):
    end_date = datetime.now().strftime("%Y/%m/%d")
    start_date = (datetime.now() - timedelta(days=days)).strftime("%Y/%m/%d")
    date_query = f"({start_date}[PDAT] : {end_date}[PDAT])"

    terms = [f'({w.strip()}[Title/Abstract])' for w in keywords.split() if w.strip()]
    query = " AND ".join(terms) + f" AND {date_query}"

    handle = Entrez.esearch(db="pubmed", term=query, retmax=50)
    record = Entrez.read(handle)
    ids = record["IdList"]

    if not ids:
        return []

    handle = Entrez.efetch(db="pubmed", id=ids, rettype="medline", retmode="text")
    articles = list(Medline.parse(handle))
    results = []
    for art in articles:
        results.append({
            "source": "PubMed",
            "title": art.get("TI", ""),
            "authors": ", ".join(art.get("AU", [])),
            "journal": art.get("JT", ""),
            "year": art.get("DP", "").split(" ")[0],
            "doi": art.get("LID", "").replace(" [doi]", ""),
            "link": f"https://pubmed.ncbi.nlm.nih.gov/{art.get('PMID', '')}/"
        })
    return results

# === Форматування у стилі Vancouver ===
def format_vancouver(entry):
    return (f"{entry['authors']}. {entry['title']}. {entry['journal']}. "
            f"{entry['year']}. doi:{entry['doi']}\n{entry['link']}")

# === Збереження у файли ===
def save_results(results):
    df = pd.DataFrame(results)
    csv_file = "results.csv"
    df.to_csv(csv_file, index=False)

    doc = Document()
    doc.add_heading("PubMed Search Results (Vancouver Style)", level=1)
    for i, r in enumerate(results, 1):
        doc.add_paragraph(f"{i}. {format_vancouver(r)}")
    docx_file = "results.docx"
    doc.save(docx_file)

    return csv_file, docx_file

# === Команда /start ===
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "👋 Привіт! Надішли мені ключові слова англійською (наприклад: 'gpt urology ai')."
    )
    return ASK_KEYWORDS

# === Обробка ключових слів ===
async def get_keywords(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data["keywords"] = update.message.text.strip()
    await update.message.reply_text("📅 За скільки останніх днів брати статті? (натисни Enter для 30 днів)")
    return ASK_DAYS

# === Обробка кількості днів ===
async def get_days(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = update.message.text.strip()
    days = int(text) if text.isdigit() else 30
    keywords = context.user_data["keywords"]

    await update.message.reply_text(f"🔎 Шукаю статті за запитом: {keywords} (останні {days} днів)...")

    results = fetch_pubmed_full(keywords, days)

    if not results:
        await update.message.reply_text("❌ Нічого не знайдено. Спробуй інші ключові слова.")
    else:
        csv_file, docx_file = save_results(results)
        await update.message.reply_text(f"✅ Знайдено {len(results)} статей.")
        await update.message.reply_document(open(csv_file, "rb"))
        await update.message.reply_document(open(docx_file, "rb"))

    await update.message.reply_text("🔁 Надішли нові ключові слова або /start для початку.")
    return ASK_KEYWORDS

# === Основна функція ===
def main():
    TOKEN = os.getenv("BOT_TOKEN")
    if not TOKEN:
        raise ValueError("❌ BOT_TOKEN не знайдено! Додай його у Railway → Variables")

    app = ApplicationBuilder().token(TOKEN).build()

    conv_handler = ConversationHandler(
        entry_points=[CommandHandler("start", start)],
        states={
            ASK_KEYWORDS: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_keywords)],
            ASK_DAYS: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_days)],
        },
        fallbacks=[CommandHandler("start", start)],
    )

    app.add_handler(conv_handler)

    # === Автоматичне налаштування режиму запуску ===
    port = int(os.environ.get("PORT", 8080))
    url = os.environ.get("RAILWAY_STATIC_URL")

    if url:
        webhook_url = f"https://{url}/webhook"
        print(f"✅ Встановлюю webhook: {webhook_url}")
        app.run_webhook(
            listen="0.0.0.0",
            port=port,
            url_path="/webhook",
            webhook_url=webhook_url
        )
    else:
        print("⚙️ Webhook не знайдено, запуск локально через polling...")
        app.run_polling()

# === Точка входу ===
if __name__ == "__main__":
    main()
